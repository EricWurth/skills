#!/usr/bin/env python3
"""Render a resumebot markdown resume (master cut / variant) to a .docx.

Consistent professional format: Calibri, centered 16pt bold name, 11pt bold
section headers, 10pt body, 0.5" margins. Requires python-docx.

Input markdown conventions (see examples/ExampleMasterResume.md):
  # Name                     -> centered 16pt bold
  first non-empty line after -> centered contact line (pipes kept)
  ## Section                 -> bold 11pt section header with rule line
  ### Employer — Location    -> bold subheading
  **Title** | dates          -> italic-free bold line (role line)
  - bullet                   -> bulleted 10pt
  plain paragraph            -> 10pt body
Blockquote lines (>) and horizontal rules (---) are skipped: they are
master-document annotations, never submission content.

Usage:
  python build_resume.py input.md output.docx [--name-size 16]
"""

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

BODY = 10
HEADER = 11


def add_run(para, text, bold=False, size=BODY):
    """Add text, honoring **bold** spans (structural bold only — the lint
    gate rejects mid-sentence bold in generated output, so use sparingly)."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = para.add_run(part[2:-2])
            r.bold = True
        else:
            r = para.add_run(part)
            r.bold = bold
        r.font.size = Pt(size)
        r.font.name = "Calibri"


def build(md_path: Path, out_path: Path, name_size: int):
    doc = Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Inches(0.5)
        section.left_margin = section.right_margin = Inches(0.5)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(BODY)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    saw_name = False
    expect_contact = False

    for raw in lines:
        line = raw.rstrip()
        stripped = line.strip()
        if not stripped or stripped.startswith(">") or set(stripped) <= {"-"} and len(stripped) >= 3:
            continue

        if stripped.startswith("# ") and not saw_name:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, stripped[2:].replace("**", ""), bold=True, size=name_size)
            p.paragraph_format.space_after = Pt(2)
            saw_name = True
            expect_contact = True
        elif expect_contact and not stripped.startswith("#"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", stripped).replace("**", ""), size=BODY)
            p.paragraph_format.space_after = Pt(8)
            expect_contact = False
        elif stripped.startswith("## "):
            p = doc.add_paragraph()
            add_run(p, stripped[3:].upper(), bold=True, size=HEADER)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(3)
        elif stripped.startswith("### "):
            p = doc.add_paragraph()
            add_run(p, stripped[4:], bold=True, size=BODY)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(1)
        elif stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_run(p, stripped[2:], size=BODY)
            p.paragraph_format.space_after = Pt(1)
        else:
            p = doc.add_paragraph()
            add_run(p, stripped, size=BODY)
            p.paragraph_format.space_after = Pt(3)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"wrote {out_path}")


def main():
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("input")
    ap.add_argument("output")
    ap.add_argument("--name-size", type=int, default=16)
    args = ap.parse_args()
    if not Path(args.input).exists():
        sys.exit(f"error: {args.input} not found")
    build(Path(args.input), Path(args.output), args.name_size)


if __name__ == "__main__":
    main()
